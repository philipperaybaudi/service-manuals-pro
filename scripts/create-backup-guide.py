"""
Génère le guide de sauvegarde en Word sur le Bureau.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5) if p.runs else None
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    p._p.get_or_add_pPr().append(shd)
    return p

def spacer():
    doc.add_paragraph()

# ── Titre principal ───────────────────────────────────────────────────────────
titre = doc.add_heading('Guide de Sauvegarde — Service Manuals Pro', 0)
titre.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(f'Généré le {datetime.date.today().strftime("%d/%m/%Y")}')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
sub.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
h1('1. Ce que contient le dossier Claude Doc GB test')
# ══════════════════════════════════════════════════════════════════════════════

body('Le dossier C:\\Users\\adm\\Claude Doc GB test pèse environ 2 Go mais la majorité est du cache inutile. Voici la décomposition :')
spacer()

bullet('service-manuals-pro/src → ', bold_prefix='Code source du site (Next.js)')
bullet('service-manuals-pro/scripts → ', bold_prefix='Tous les scripts du pipeline import')
bullet('service-manuals-pro/.env.local → ', bold_prefix='Clés secrètes (Supabase, R2, Stripe, Anthropic) — CRITIQUE')
bullet('service-manuals-pro/CLAUDE.md → ', bold_prefix='Règles du pipeline sécurisé')
bullet('node_modules/ (1 Go) → ', bold_prefix='À NE PAS sauvegarder — ')
bullet('node_modules/ (.next, .wrangler) → ', bold_prefix='À NE PAS sauvegarder — ')

p = doc.add_paragraph()
run = p.add_run('Important : ')
run.bold = True
run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
p.add_run('Ce dossier ne contient NI les PDFs (sur R2 Cloudflare), NI la base de données (sur Supabase), NI les previews (sur Supabase Storage). Une sauvegarde complète nécessite donc 3 sources distinctes.')

spacer()

# ══════════════════════════════════════════════════════════════════════════════
h1('2. Dossiers C: à sauvegarder')
# ══════════════════════════════════════════════════════════════════════════════

body('Sauvegarder uniquement ces dossiers/fichiers (fréquence : hebdomadaire) :')
spacer()

h2('2.1 Code et scripts (CRITIQUE)')
items = [
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\src', 'Code source du site'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\scripts', 'Pipeline import complet'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\.env.local', 'Clés secrètes — à stocker en lieu sûr'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\CLAUDE.md', 'Règles pipeline sécurisé'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\wrangler.toml', 'Config Cloudflare'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\next.config.mjs', 'Config Next.js'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\package.json', 'Dépendances npm'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\tailwind.config.ts', 'Config CSS'),
    ('C:\\Users\\adm\\Claude Doc GB test\\service-manuals-pro\\tsconfig.json', 'Config TypeScript'),
]
for path, desc in items:
    bullet(f'{path}\n   → {desc}')

spacer()
h2('2.2 Archive locale des PDFs (CRITIQUE)')
bullet('C:\\Users\\adm\\Documents\\SHEMATHEQUE → Archive complète : DOSSIER SOURCE, DOCS EN LIGNE, et tous les PDFs traités')

spacer()
h2('2.3 Mémoire Claude (Utile)')
bullet('C:\\Users\\adm\\.claude\\projects → Mémoire des sessions Claude (MEMORY.md, règles, contexte projet)')

spacer()

# ══════════════════════════════════════════════════════════════════════════════
h1('3. Sauvegarder Supabase (base de données + previews)')
# ══════════════════════════════════════════════════════════════════════════════

h2('3.1 Base de données PostgreSQL')
body('Option A — Via le dashboard Supabase (sans installation) :')
bullet('Aller sur supabase.com → ton projet → Settings → Database → Backups')
bullet('Télécharger le dernier backup automatique (plans payants)')
bullet("Ou : Table Editor → documents → Export CSV (répéter pour 'brands' et 'categories')")

spacer()
body('Option B — Via pg_dump en ligne de commande :')
body('Les identifiants de connexion se trouvent dans Supabase → Settings → Database → Connection string.')
code('pg_dump "postgresql://postgres:[MOT_DE_PASSE]@[HOST].supabase.co:5432/postgres" --table=documents --table=brands --table=categories -F c -f "D:\\BACKUP\\supabase-dump-2026-05-19.dump"')

spacer()
h2('3.2 Storage Supabase (previews JPG)')
body('Le bucket "logos" contient toutes les previews (previews/{slug}.jpg). Via la Supabase CLI :')
code('npm install -g supabase')
code('supabase login')
code('supabase storage cp --recursive "ss:///logos/previews" "D:\\BACKUP\\supabase-previews\\"')

p = doc.add_paragraph()
p.add_run('Note : ').bold = True
p.add_run('les previews sont régénérables depuis les PDFs source. Priorité moindre que la base de données.')

spacer()

# ══════════════════════════════════════════════════════════════════════════════
h1('4. Sauvegarder R2 Cloudflare (PDFs)')
# ══════════════════════════════════════════════════════════════════════════════

body('R2 est compatible S3. L\'outil recommandé est rclone (gratuit, open-source).')

h2('4.1 Installation rclone (une seule fois)')
bullet('Télécharger sur rclone.org/downloads/')
bullet('Dézipper → placer rclone.exe dans un dossier accessible (ex: C:\\Tools\\)')

spacer()
h2('4.2 Configuration rclone')
code('rclone config')
body('Répondre aux questions :')
bullet('n → nouveau remote')
bullet('Nom : r2')
bullet('Storage : s3 (Cloudflare R2 est compatible S3)')
bullet('Provider : Cloudflare')
bullet('Access Key ID : valeur R2_ACCESS_KEY_ID dans .env.local')
bullet('Secret Access Key : valeur R2_SECRET_ACCESS_KEY dans .env.local')
bullet('Endpoint : https://[R2_ACCOUNT_ID].r2.cloudflarestorage.com')

spacer()
h2('4.3 Synchronisation vers DD externe ou NAS')
body('Vers DD externe :')
code('rclone sync r2:service-manuals-documents "D:\\BACKUP\\r2-pdfs\\" --progress')
spacer()
body('Vers NAS Synology :')
code('rclone sync r2:service-manuals-documents "\\\\NAS\\backup\\r2-pdfs\\" --progress')

p = doc.add_paragraph()
p.add_run('Fréquence recommandée : ').bold = True
p.add_run('mensuelle, ou après chaque lot d\'import important. rclone ne télécharge que les fichiers nouveaux/modifiés.')

spacer()

# ══════════════════════════════════════════════════════════════════════════════
h1('5. Récupération après plantage')
# ══════════════════════════════════════════════════════════════════════════════

h2('5.1 Plantage PC')
body('Avec les sauvegardes en place, le site est restaurable en moins d\'une heure :')
bullet('1. Copier le dossier service-manuals-pro sur le nouveau PC')
bullet('2. npm install → recrée node_modules')
bullet('3. Vérifier .env.local → reconnecte Supabase + R2 + Stripe')
bullet('4. npm run build && npx wrangler pages deploy → redéploie le site')
bullet('5. Les PDFs (R2) et la base de données (Supabase) sont intacts — rien à restaurer')

spacer()
h2('5.2 Plantage Supabase')
body('Restaurer depuis le dump PostgreSQL :')
code('pg_restore -d "postgresql://postgres:[MOT_DE_PASSE]@[HOST]:5432/postgres" "D:\\BACKUP\\supabase-dump-XXXX.dump"')

spacer()
h2('5.3 Plantage R2')
body('Restaurer les PDFs depuis la sauvegarde locale :')
code('rclone sync "D:\\BACKUP\\r2-pdfs\\" r2:service-manuals-documents --progress')

spacer()

# ══════════════════════════════════════════════════════════════════════════════
h1('6. Tableau récapitulatif')
# ══════════════════════════════════════════════════════════════════════════════

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# En-tête
hdr = table.rows[0].cells
hdr[0].text = 'Quoi'
hdr[1].text = 'Fréquence'
hdr[2].text = 'Durée estimée'
hdr[3].text = 'Priorité'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E3A8A')
    cell._tc.get_or_add_tcPr().append(shd)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows = [
    ('Dossiers C: (code + scripts + .env.local)', 'Hebdomadaire', '< 5 minutes', '🔴 Critique'),
    ('SHEMATHEQUE (archive PDFs locaux)', 'Hebdomadaire', '10-30 min', '🔴 Critique'),
    ('Export CSV Supabase (documents, brands)', 'Après chaque lot importé', '2 minutes', '🔴 Critique'),
    ('Sync R2 → DD externe (rclone)', 'Mensuelle', '1-3 heures', '🟠 Important'),
    ('Previews Supabase Storage', 'Mensuelle', 'Variable', '🟡 Utile (régénérables)'),
    ('Mémoire Claude (.claude/projects)', 'Mensuelle', '< 1 minute', '🟡 Utile'),
]
for row_data in rows:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

spacer()

# ── Pied de page ──────────────────────────────────────────────────────────────
p = doc.add_paragraph('Document généré automatiquement par Claude Code — Service Manuals Pro')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
p.runs[0].font.size = Pt(9)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out = r'C:\Users\adm\Desktop\Guide-Sauvegarde-ServiceManualsPro.docx'
doc.save(out)
print(f'Document créé : {out}')
